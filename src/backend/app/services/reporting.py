from io import BytesIO
from pathlib import Path
import subprocess

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.reporting import ReportRequest


def _safe_name(value: str) -> str:
    return "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in value).strip("_")


def render_report_markdown(payload: ReportRequest, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    file_name = f"{_safe_name(payload.project_name)}_report.md"
    report_path = output_dir / file_name

    ranking_rows = "\n".join(
        [
            f"| {r.option_name} | {r.risk_adjusted_score:.2f} | {r.expected_score:.2f} | {'Pass' if r.passed_gates else 'Fail'} |"
            for r in payload.ranking
        ]
    )
    assumptions_block = "\n".join([f"- {item}" for item in payload.assumptions])

    # Criteria score summary per option
    criteria_summary_blocks = []
    for r in payload.ranking:
        gate_label = "PASS" if r.passed_gates else f"FAIL ({', '.join(r.gate_failures)})"
        rows = "\n".join(
            [
                f"| {cat} | {score:.2f} |"
                for cat, score in sorted(r.category_scores.items())
            ]
        )
        block = f"""### {r.option_name} (Gate: {gate_label})

| Category | Mean Score (0–10) |
|---|---:|
{rows}
"""
        criteria_summary_blocks.append(block)

    criteria_summary_section = "\n".join(criteria_summary_blocks) if criteria_summary_blocks else "_No criteria scores available._"

    text = f"""---
title: \"{payload.project_name} - Weighted SWOT Report\"
format:
  pdf: default
  docx: default
---

# Executive Summary

{payload.executive_summary}

# Project Scope

{payload.project_scope or '_No scope provided._'}

# Decision Outcome

- Winner: {payload.winner or 'None'}
- Loser: {payload.loser or 'None'}

# Option Ranking

| Option | Risk-Adjusted Score | Expected Score | Gate |
|---|---:|---:|---|
{ranking_rows}

# Criteria Score Summary by Option

{criteria_summary_section}

# Assumptions

{assumptions_block}
"""

    report_path.write_text(text, encoding="utf-8")
    return report_path


def render_report_docx(payload: ReportRequest) -> bytes:
    """Build a DOCX in memory using python-docx and return raw bytes."""
    doc = Document()

    # Title
    title = doc.add_heading(f"{payload.project_name} — Weighted SWOT Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if payload.client_name:
        client_para = doc.add_paragraph(f"Client: {payload.client_name}")
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(payload.executive_summary)

    doc.add_heading("Project Scope", level=1)
    doc.add_paragraph(payload.project_scope or "No scope provided.")

    doc.add_heading("Decision Outcome", level=1)
    doc.add_paragraph(f"Winner: {payload.winner or 'None'}")
    doc.add_paragraph(f"Loser: {payload.loser or 'None'}")

    doc.add_heading("Option Ranking", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Option"
    hdr[1].text = "Risk-Adjusted Score"
    hdr[2].text = "Expected Score"
    hdr[3].text = "Gate"
    for cell in hdr:
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.bold = True

    for r in payload.ranking:
        row = table.add_row().cells
        row[0].text = r.option_name
        row[1].text = f"{r.risk_adjusted_score:.2f}"
        row[2].text = f"{r.expected_score:.2f}"
        row[3].text = "Pass" if r.passed_gates else "Fail"

    doc.add_heading("Assumptions", level=1)
    for assumption in payload.assumptions:
        doc.add_paragraph(assumption, style="List Bullet")

    doc.add_heading("Criteria Score Summary by Option", level=1)
    for r in payload.ranking:
        gate_label = "PASS" if r.passed_gates else f"FAIL ({', '.join(r.gate_failures)})"
        doc.add_heading(f"{r.option_name}  —  Gate: {gate_label}", level=2)
        if r.category_scores:
            cat_table = doc.add_table(rows=1, cols=2)
            cat_table.style = "Table Grid"
            cat_hdr = cat_table.rows[0].cells
            cat_hdr[0].text = "Category"
            cat_hdr[1].text = "Mean Score (0–10)"
            for cell in cat_hdr:
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
                run.bold = True
            for cat, score in sorted(r.category_scores.items()):
                cat_row = cat_table.add_row().cells
                cat_row[0].text = cat
                cat_row[1].text = f"{score:.2f}"
        else:
            doc.add_paragraph("No category scores recorded.", style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_latex_placeholder(markdown_path: Path) -> Path:
    latex_path = markdown_path.with_suffix(".tex")
    latex_path.write_text(
        "% Placeholder LaTeX output. Use Quarto/Pandoc to render polished PDF in Phase 1.1\n"
        + "% Source markdown: "
        + str(markdown_path)
        + "\n",
        encoding="utf-8",
    )
    return latex_path


def try_render_with_quarto(markdown_path: Path, dotx_template: Path) -> dict:
    pdf_path = markdown_path.with_suffix(".pdf")
    docx_path = markdown_path.with_suffix(".docx")

    try:
        subprocess.run(
            ["quarto", "render", str(markdown_path), "--to", "pdf"],
            check=True,
            capture_output=True,
            text=True,
        )
        subprocess.run(
            [
                "quarto",
                "render",
                str(markdown_path),
                "--to",
                "docx",
                "--output",
                str(docx_path.name),
                "--metadata",
                f"reference-doc={dotx_template}",
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        return {
            "status": "rendered",
            "pdf_path": str(pdf_path if pdf_path.exists() else markdown_path.with_suffix(".pdf")),
            "docx_path": str(docx_path if docx_path.exists() else markdown_path.with_suffix(".docx")),
        }
    except (subprocess.CalledProcessError, FileNotFoundError):
        return {
            "status": "quarto-unavailable",
            "pdf_path": None,
            "docx_path": None,
        }
